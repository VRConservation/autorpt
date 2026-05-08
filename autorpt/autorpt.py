#!/usr/bin/env python
"""
Auto-report generator from content.md with budget table insertion.
Reads content.md and replaces [insert budget from budget.xlsx here] with actual budget table.
"""

import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import re
import argparse
import subprocess


def read_excel_as_dataframe(excel_file):
    """Read Excel file and return pandas DataFrame"""
    try:
        df = pd.read_excel(excel_file)
        return df
    except Exception as e:
        print(f"❌ Error reading Excel file: {e}")
        return None


def add_table_to_document(doc, df):
    """Add a pandas DataFrame as a table to Word document"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    # Create table with header row + data rows
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    
    # Remove all shading and borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        tbl.insert(0, tblPr)
    
    # Remove existing borders if any
    tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    
    # Add new borders with all set to none
    no_border_xml = f'<w:tblBorders {nsdecls("w")}><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tblBorders>'
    borders_element = parse_xml(no_border_xml)
    tblPr.append(borders_element)
    
    # Add header row (bolded)
    header_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        header_cells[i].text = str(column)
        # Make header bold
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data rows
    for row_idx, (_, row) in enumerate(df.iterrows(), start=1):
        data_cells = table.rows[row_idx].cells
        is_total_row = str(row.iloc[0]).upper() == 'TOTAL'
        
        for col_idx, value in enumerate(row):
            # Handle NaN and None values
            if pd.isna(value):
                data_cells[col_idx].text = ""
            # Format numeric values without $ signs
            elif isinstance(value, (int, float)):
                if col_idx > 0:  # Only for numeric columns after first
                    data_cells[col_idx].text = f"{value:,.0f}" if value == int(value) else f"{value:,.2f}"
                else:
                    data_cells[col_idx].text = str(value)
            else:
                data_cells[col_idx].text = str(value)
            
            # Bold only TOTAL row
            if is_total_row:
                for paragraph in data_cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True


def parse_markdown_sections(markdown_text):
    """Parse markdown text into sections by headers"""
    sections = []
    current_section = None
    current_content = []
    
    for line in markdown_text.split('\n'):
        if line.startswith('# '):
            # Top-level header
            if current_section is not None:
                sections.append({
                    'level': 1,
                    'title': current_section,
                    'content': '\n'.join(current_content).strip()
                })
            current_section = line[2:].strip()
            current_content = []
        elif line.startswith('## '):
            # Sub-header
            if current_section is not None:
                sections.append({
                    'level': 1,
                    'title': current_section,
                    'content': '\n'.join(current_content).strip()
                })
            current_section = line[3:].strip()
            current_content = []
        else:
            current_content.append(line)
    
    # Don't forget the last section
    if current_section is not None:
        sections.append({
            'level': 1,
            'title': current_section,
            'content': '\n'.join(current_content).strip()
        })
    
    return sections


def add_markdown_to_document(doc, markdown_text, budget_df=None):
    """Add markdown-formatted content to Word document"""
    lines = markdown_text.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        if not line:
            # Skip empty lines - don't add paragraph spacing
            i += 1
            continue
        elif line.startswith('## '):
            # Level 2 header
            heading = line[3:].strip()
            doc.add_heading(heading, level=2)
            i += 1
        elif line.startswith('[insert budget from budget.xlsx here]'):
            # Replace with actual budget table
            if budget_df is not None:
                add_table_to_document(doc, budget_df)
            i += 1
        elif line.startswith('- '):
            # Bullet point - collect all consecutive bullet points
            bullets = []
            while i < len(lines) and lines[i].rstrip().startswith('- '):
                bullet_text = lines[i].rstrip()[2:].strip()
                bullets.append(bullet_text)
                i += 1
            # Add all bullets together
            for bullet_text in bullets:
                doc.add_paragraph(bullet_text, style='List Bullet')
        else:
            # Regular paragraph
            if line.strip():
                doc.add_paragraph(line)
            i += 1


def generate_report_from_content():
    """Main function to generate report"""
    reports_dir = Path('reports')
    content_file = reports_dir / 'content.md'
    budget_file = reports_dir / 'budget.xlsx'
    
    print("📄 Reading content.md...")
    if not content_file.exists():
        print(f"❌ {content_file} not found")
        return False
    
    with open(content_file, 'r', encoding='utf-8') as f:
        content_text = f.read()
    
    print("📊 Reading budget.xlsx...")
    if not budget_file.exists():
        print(f"❌ {budget_file} not found")
        return False
    
    budget_df = read_excel_as_dataframe(budget_file)
    if budget_df is None:
        return False
    
    print("📝 Creating Word document...")
    doc = Document()
    
    # Parse content into sections
    sections = parse_markdown_sections(content_text)
    
    for section in sections:
        title = section['title']
        content = section['content']
        
        if title.lower() == 'monthly report':
            # Add as main title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add today's date under title
            date_para = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Add as regular section
            if title:
                doc.add_heading(title, level=1)
            
            # Add content
            if content:
                add_markdown_to_document(doc, content, budget_df)
    
    # Save document
    output_file = reports_dir / f"test_report_{datetime.now().strftime('%Y-%m-%d')}.docx"
    doc.save(str(output_file))
    print(f"✅ Report generated: {output_file}")
    return True


def df_to_typst_table(df):
    """Convert a DataFrame to Typst table syntax with formatting"""
    cells = []
    
    # Add header row with bold formatting
    for col in df.columns:
        cells.append(f'[*{col}*]')
    
    # Add data rows
    for idx, (_, row) in enumerate(df.iterrows()):
        task_name = str(row.iloc[0]).strip() if not pd.isna(row.iloc[0]) else ''
        
        # Determine row type
        is_total_project = 'total project' in task_name.lower()
        is_indirect = 'indirect' in task_name.lower()
        is_sub_total = 'subtotal' in task_name.lower() or 'sub-total' in task_name.lower() or 'sub total' in task_name.lower()
        is_total_row = 'total' in task_name.lower() and not is_total_project and not is_indirect and not is_sub_total
        
        # Check if this is a sub-item: has decimal number (2.1, 2.2, etc) or no number following a numbered category
        is_sub_item = False
        
        # Check if current task has decimal number (2.1, 2.2, etc.)
        import re
        has_decimal_number = bool(re.match(r'^\d+\.\d+', task_name))
        
        if has_decimal_number and not is_total_project and not is_indirect and not is_total_row and not is_sub_total:
            is_sub_item = True
        elif idx > 0:
            prev_task = str(df.iloc[idx-1, 0]).strip() if not pd.isna(df.iloc[idx-1, 0]) else ''
            # Previous row is a category if it starts with a number (1., 2., 3., etc.)
            prev_is_category = any(f'{i}.' in prev_task for i in range(1, 10)) and not bool(re.match(r'^\d+\.\d+', prev_task))
            # Current row is a sub-item if it doesn't start with a number
            curr_is_category = any(f'{i}.' in task_name for i in range(1, 10))
            
            if prev_is_category and not curr_is_category and not is_total_project and not is_indirect and not is_total_row and not is_sub_total:
                is_sub_item = True
        
        # Add horizontal line above Total Project
        if is_total_project:
            cells.append('table.hline()')
        
        # Format task name
        if is_total_row:
            formatted_task = f'*{task_name}*'
        elif is_total_project:
            # Capitalize Total Project
            formatted_task = task_name.replace('total project', 'Total Project').replace('TOTAL PROJECT', 'Total Project')
        elif is_indirect:
            # Capitalize Indirect
            formatted_task = task_name.replace('indirect', 'Indirect').replace('INDIRECT', 'Indirect')
        elif is_sub_total:
            formatted_task = task_name
        elif is_sub_item:
            formatted_task = f'    {task_name}'  # Indent with 4 spaces
        else:
            formatted_task = task_name
        
        cells.append(f'[{formatted_task}]')
        
        # Add numeric columns
        for col_idx in range(1, len(row)):
            val = row.iloc[col_idx]
            
            # Convert to string, handling NaN values
            if pd.isna(val):
                # For category headers without values, leave blank
                if any(f'{i}.' in task_name for i in range(1, 10)):
                    cell_str = ''
                else:
                    cell_str = '-'
            else:
                # Format numbers with comma separator and no decimals
                try:
                    num = float(val)
                    cell_str = f'{num:,.0f}'
                except (ValueError, TypeError):
                    cell_str = str(val)
            
            # Bold only the TOTAL row (not subtotal or total project/indirect numbers)
            if is_total_row:
                cells.append(f'[#align(right)[*{cell_str}*]]')
            elif is_total_project or is_indirect:
                cells.append(f'[#align(right)[*{cell_str}*]]')
            else:
                cells.append(f'[#align(right)[{cell_str}]]')
        
        # Add horizontal line below Indirect
        if is_indirect:
            cells.append('table.hline()')
    
    # Format as Typst table
    table_str = f'#table(\n  columns: {len(df.columns)},\n  ' + ', '.join(cells) + '\n)'
    
    return table_str


def generate_pdf_with_typst():
    """Generate PDF from content.md using Typst"""
    reports_dir = Path('reports')
    content_file = reports_dir / 'content.md'
    budget_file = reports_dir / 'budget.xlsx'
    report_content_file = reports_dir / 'report_content.typ'
    typst_template = reports_dir / 'report.typ'
    
    print("📄 Checking Typst template and content.md...")
    
    if not typst_template.exists():
        print(f"❌ Typst template not found: {typst_template}")
        print("   Please create report.typ in the reports/ folder")
        return False
    
    if not content_file.exists():
        print(f"❌ Content file not found: {content_file}")
        return False
    
    # Read budget if it exists
    budget_df = None
    if budget_file.exists():
        budget_df = read_excel_as_dataframe(budget_file)
    
    # Read and parse content.md
    with open(content_file, 'r', encoding='utf-8') as f:
        content_text = f.read()
    
    # Strip YAML frontmatter
    lines = content_text.split('\n')
    body_start = 0
    if lines[0].strip() == '---':
        for i in range(1, len(lines)):
            if lines[i].strip() == '---':
                body_start = i + 1
                break
    
    body_lines = lines[body_start:]
    
    # Convert markdown to Typst format
    typst_content = []
    i = 0
    while i < len(body_lines):
        line = body_lines[i].rstrip()
        
        if not line:
            i += 1
            continue
        elif line.startswith('# '):
            heading = line[2:].strip()
            typst_content.append(f'= {heading}')
            typst_content.append('')
            i += 1
        elif line.startswith('## '):
            heading = line[3:].strip()
            typst_content.append(f'== {heading}')
            typst_content.append('')
            i += 1
        elif line.startswith('[insert budget from budget.xlsx here]'):
            # Insert actual budget table
            if budget_df is not None and not budget_df.empty:
                typst_content.append(df_to_typst_table(budget_df))
                typst_content.append('')
            i += 1
        elif line.startswith('- '):
            # Bullet points
            while i < len(body_lines) and body_lines[i].rstrip().startswith('- '):
                bullet = body_lines[i].rstrip()[2:].strip()
                typst_content.append(f'- {bullet}')
                i += 1
            typst_content.append('')
        else:
            typst_content.append(line)
            i += 1
    
    # Write report_content.typ
    with open(report_content_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(typst_content).strip())
    
    # Generate PDF filename
    output_file = reports_dir / f"report_{datetime.now().strftime('%Y-%m-%d')}.pdf"
    
    try:
        print(f"🔄 Converting to PDF with Typst...")
        # Run typst compile command
        result = subprocess.run(
            ['typst', 'compile', 'report.typ', f"report_{datetime.now().strftime('%Y-%m-%d')}.pdf"],
            cwd=str(reports_dir),
            capture_output=True,
            text=True,
            timeout=30
        )
        
        if result.returncode == 0:
            if output_file.exists():
                file_size = output_file.stat().st_size / 1024  # KB
                print(f"✅ PDF created successfully: {output_file}")
                print(f"📊 PDF file size: {file_size:.1f} KB")
                return True
            else:
                print("❌ Typst compilation succeeded but PDF file not created")
                return False
        else:
            print(f"❌ Typst compilation failed:")
            print(result.stderr)
            return False
    
    except FileNotFoundError:
        print("❌ Typst not found. Please install Typst:")
        print("   https://github.com/typst/typst")
        return False
    except subprocess.TimeoutExpired:
        print("❌ Typst compilation timed out")
        return False
    except Exception as e:
        print(f"❌ Error generating PDF: {e}")
        return False



    """Main function to generate report"""
    reports_dir = Path('reports')
    content_file = reports_dir / 'content.md'
    budget_file = reports_dir / 'budget.xlsx'
    
    print("📄 Reading content.md...")
    if not content_file.exists():
        print(f"❌ {content_file} not found")
        return False
    
    with open(content_file, 'r', encoding='utf-8') as f:
        content_text = f.read()
    
    print("📊 Reading budget.xlsx...")
    if not budget_file.exists():
        print(f"❌ {budget_file} not found")
        return False
    
    budget_df = read_excel_as_dataframe(budget_file)
    if budget_df is None:
        return False
    
    print("📝 Creating Word document...")
    doc = Document()
    
    # Parse content into sections
    sections = parse_markdown_sections(content_text)
    
    for section in sections:
        title = section['title']
        content = section['content']
        
        if title.lower() == 'monthly report':
            # Add as main title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add today's date under title
            date_para = doc.add_paragraph(datetime.now().strftime('%B %d, %Y'))
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Add as regular section
            if title:
                doc.add_heading(title, level=1)
            
            # Add content
            if content:
                add_markdown_to_document(doc, content, budget_df)
    
    # Save document
    output_file = reports_dir / f"test_report_{datetime.now().strftime('%Y-%m-%d')}.docx"
    doc.save(str(output_file))
    print(f"✅ Report generated: {output_file}")
    return True


def main():
    """Entry point for console script"""
    parser = argparse.ArgumentParser(
        description='Generate automated reports from content.md',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  auto                          # Generate Word report (.docx)
  auto --typst                  # Generate PDF report using Typst
  auto --pdf                    # Generate both Word and PDF
        """)
    
    parser.add_argument('--typst', action='store_true',
                        help='Generate PDF using Typst instead of .docx')
    parser.add_argument('--pdf', action='store_true',
                        help='Generate both .docx and PDF (Typst)')
    parser.add_argument('--verbose', '-v', action='store_true',
                        help='Enable verbose output')
    
    args = parser.parse_args()
    
    if args.verbose:
        print("🔧 Verbose mode enabled")
    
    # If typst or pdf flags are set, use typst for PDF generation
    if args.typst:
        print("🚀 Generating PDF report with Typst...")
        success = generate_pdf_with_typst()
        if success:
            print("\n✅ PDF generation completed successfully!")
            return 0
        else:
            print("\n❌ PDF generation failed")
            return 1
    elif args.pdf:
        print("🚀 Generating Word and PDF reports...")
        docx_success = generate_report_from_content()
        pdf_success = generate_pdf_with_typst()
        if docx_success and pdf_success:
            print("\n✅ Report generation completed successfully!")
            return 0
        else:
            print("\n❌ Some reports failed to generate")
            return 1
    else:
        # Default: generate Word report
        print("🚀 Starting auto-report generation...")
        success = generate_report_from_content()
        
        if success:
            print("\n✅ Report generation completed successfully!")
            return 0
        else:
            print("\n❌ Report generation failed")
            return 1


if __name__ == '__main__':
    exit(main())
